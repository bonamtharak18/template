import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class SimpleResumeGenerator:
    def __init__(self):
        self.document = docx.Document()
        self.font_name = "Arial"
        self.font_size = Pt(11)

    def set_default_styles(self):
        """Sets some basic default styles."""
        style = self.document.styles['Normal']
        font = style.font
        font.name = self.font_name
        font.size = self.font_size

    def add_header(self, name, email, phone):
        """Adds the header section."""
        self.set_default_styles()  # Important to apply default styles
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align
        run = paragraph.add_run(name)
        font = run.font
        font.name = self.font_name
        font.size = Pt(16)  # Larger font for name
        font.bold = True

        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"{email} | {phone}")
        font = run.font
        font.name = self.font_name
        font.size = self.font_size


    def add_section_heading(self, title):
        """Adds a section heading."""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(title)
        font = run.font
        font.name = self.font_name
        font.size = Pt(14)
        font.bold = True
        paragraph.space_before = Pt(12)  # Add space before heading


    def add_bullet_points(self, items):
        """Adds bullet points."""
        for item in items:
            paragraph = self.document.add_paragraph(item, style='List Bullet')
            paragraph.paragraph_format.left_indent = Inches(0.25)  # Add an indent



    def add_experience(self, company, title, dates, description):
        """Adds an experience entry."""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(f"{title}, {company} ({dates})")
        font = run.font
        font.name = self.font_name
        font.size = self.font_size
        font.bold = True

        self.add_bullet_points(description)

    def add_education(self, institution, degree, dates):
        """Adds an education entry."""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(f"{degree}, {institution} ({dates})")
        font = run.font
        font.name = self.font_name
        font.size = self.font_size
        font.bold = True


    def save(self, filename="resume.docx"):
        """Saves the document."""
        self.document.save(filename)


# Example Usage:
if __name__ == '__main__':
    generator = SimpleResumeGenerator()

    # Data
    name = "Your Name"
    email = "your.email@example.com"
    phone = "123-456-7890"

    experience = [
        {
            "company": "Company A",
            "title": "Software Engineer",
            "dates": "2020-2023",
            "description": [
                "Developed web applications.",
                "Worked with Python and Django.",
                "Contributed to team projects."
            ]
        }
    ]

    education = [
        {
            "institution": "University X",
            "degree": "Bachelor of Science",
            "dates": "2016-2020"
        }
    ]

    skills = ["Python", "Django", "JavaScript", "HTML", "CSS"]

    # Generate Resume
    generator.add_header(name, email, phone)
    generator.add_section_heading("Summary")
    generator.add_bullet_points(["Enthusiastic and skilled professional."])

    generator.add_section_heading("Experience")
    for exp in experience:
        generator.add_experience(exp['company'], exp['title'], exp['dates'], exp['description'])

    generator.add_section_heading("Education")
    for edu in education:
        generator.add_education(edu['institution'], edu['degree'], edu['dates'])

    generator.add_section_heading("Skills")
    generator.add_bullet_points(skills)

    generator.save("simple_resume.docx")
    print("Resume generated successfully!")
